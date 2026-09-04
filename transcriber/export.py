"""Turn a meeting into something you can read or share."""
from __future__ import annotations

from pathlib import Path

from . import store


def _blocks(meeting: dict) -> list[dict]:
    """Merge consecutive utterances from the same speaker into paragraphs."""
    blocks = []
    for u in meeting.get("utterances", []):
        name = store.display_name(meeting, u["speaker"])
        if blocks and blocks[-1]["name"] == name:
            blocks[-1]["text"] += " " + u["text"]
        else:
            blocks.append({"name": name, "start": u["start"], "text": u["text"]})
    return blocks


def _speaker_key(meeting: dict) -> list[str]:
    lines = []
    for label, sp in sorted(meeting.get("speakers", {}).items()):
        name = store.display_name(meeting, label)
        if sp.get("confirmed"):
            lines.append(f"{name}")
        elif sp.get("guess"):
            lines.append(f"{name} (guessed, {sp.get('confidence', 'low')} confidence)")
        else:
            lines.append(f"{name} (unidentified)")
    return lines


def to_markdown(meeting: dict) -> str:
    out = [f"# {meeting['title']}", ""]
    key = _speaker_key(meeting)
    if key:
        out.append("Speakers: " + "; ".join(key))
        out.append("")
    for b in _blocks(meeting):
        out.append(f"**{b['name']}** ({store.fmt_ts(b['start'])}): {b['text']}")
        out.append("")
    return "\n".join(out).rstrip() + "\n"


def to_text(meeting: dict) -> str:
    out = [meeting["title"], ""]
    key = _speaker_key(meeting)
    if key:
        out.append("Speakers: " + "; ".join(key))
        out.append("")
    for b in _blocks(meeting):
        out.append(f"{b['name']} ({store.fmt_ts(b['start'])})")
        out.append(b["text"])
        out.append("")
    return "\n".join(out).rstrip() + "\n"


def to_docx(meeting: dict, path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(meeting["title"], level=1)
    key = _speaker_key(meeting)
    if key:
        p = doc.add_paragraph()
        p.add_run("Speakers: ").bold = True
        p.add_run("; ".join(key))
    for b in _blocks(meeting):
        p = doc.add_paragraph()
        r = p.add_run(f"{b['name']} ")
        r.bold = True
        t = p.add_run(f"({store.fmt_ts(b['start'])})  ")
        t.font.size = Pt(9)
        p.add_run(b["text"])
    doc.save(str(path))
    return path


def write(meeting: dict, fmt: str = "md") -> Path:
    d = store.meeting_dir(meeting["id"])
    if fmt == "md":
        p = d / f"{meeting['id']}.md"
        p.write_text(to_markdown(meeting))
    elif fmt == "txt":
        p = d / f"{meeting['id']}.txt"
        p.write_text(to_text(meeting))
    elif fmt == "docx":
        p = to_docx(meeting, d / f"{meeting['id']}.docx")
    else:
        raise ValueError(f"unknown format {fmt}")
    return p


# --- summaries ---------------------------------------------------------------
def _fmt_date(ts: float | None) -> str:
    import time
    return time.strftime("%B %-d, %Y", time.localtime(ts)) if ts else ""


def summary_to_markdown(meeting: dict) -> str:
    s = meeting.get("summary") or {}
    if s.get("status") != "ready":
        raise ValueError("no summary yet")
    out = [f"# {meeting['title']} — {s.get('guide_title', 'Summary')}", ""]
    key = _speaker_key(meeting)
    meta = []
    if key:
        meta.append("Speakers: " + "; ".join(key))
    if s.get("created"):
        meta.append("Summarized " + _fmt_date(s["created"]))
    if meta:
        out += ["*" + " · ".join(meta) + "*", ""]
    if s.get("overview"):
        out += ["## Overview", "", s["overview"], ""]
    if s.get("priorities"):
        out += ["## Top priorities", ""] + [f"{i}. {p}" for i, p in enumerate(s["priorities"], 1)] + [""]
    for sec in s.get("sections", []):
        out += [f"## {sec['title']}", "", f"*{sec['question']}*", ""]
        if not sec.get("covered") and not sec.get("summary"):
            out += ["_Not discussed._", ""]
            continue
        if sec.get("summary"):
            out += [sec["summary"], ""]
        if sec.get("points"):
            out += [f"- {p}" for p in sec["points"]] + [""]
        for q in sec.get("quotes", []):
            who = f" — {q['speaker']}" if q.get("speaker") else ""
            when = f" ({q['time']})" if q.get("time") else ""
            out += [f"> “{q['text']}”{who}{when}", ""]
    if s.get("follow_ups"):
        out += ["## Follow-ups", ""] + [f"- {p}" for p in s["follow_ups"]] + [""]
    return "\n".join(out).rstrip() + "\n"


def summary_to_docx(meeting: dict, path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor

    s = meeting.get("summary") or {}
    if s.get("status") != "ready":
        raise ValueError("no summary yet")
    doc = Document()
    doc.add_heading(meeting["title"], level=0)
    sub = doc.add_paragraph()
    r = sub.add_run(s.get("guide_title", "Summary"))
    r.italic = True
    key = _speaker_key(meeting)
    if key:
        p = doc.add_paragraph()
        p.add_run("Speakers: ").bold = True
        p.add_run("; ".join(key))
    if s.get("created"):
        doc.add_paragraph("Summarized " + _fmt_date(s["created"]))

    if s.get("overview"):
        doc.add_heading("Overview", level=1)
        doc.add_paragraph(s["overview"])
    if s.get("priorities"):
        doc.add_heading("Top priorities", level=1)
        for p in s["priorities"]:
            doc.add_paragraph(p, style="List Number")
    for sec in s.get("sections", []):
        doc.add_heading(sec["title"], level=1)
        q = doc.add_paragraph()
        qr = q.add_run(sec["question"])
        qr.italic = True
        qr.font.color.rgb = RGBColor(0x66, 0x6E, 0x72)
        if not sec.get("covered") and not sec.get("summary"):
            doc.add_paragraph("Not discussed.").runs[0].italic = True
            continue
        if sec.get("summary"):
            doc.add_paragraph(sec["summary"])
        for p in sec.get("points", []):
            doc.add_paragraph(p, style="List Bullet")
        for quote in sec.get("quotes", []):
            para = doc.add_paragraph(style="Intense Quote")
            para.add_run(f"“{quote['text']}”")
            tail = ""
            if quote.get("speaker"):
                tail += f" — {quote['speaker']}"
            if quote.get("time"):
                tail += f" ({quote['time']})"
            if tail:
                t = para.add_run(tail)
                t.font.size = Pt(9)
    if s.get("follow_ups"):
        doc.add_heading("Follow-ups", level=1)
        for p in s["follow_ups"]:
            doc.add_paragraph(p, style="List Bullet")
    doc.save(str(path))
    return path


def write_summary(meeting: dict, fmt: str = "md") -> Path:
    d = store.meeting_dir(meeting["id"])
    if fmt == "md":
        p = d / f"{meeting['id']}-summary.md"
        p.write_text(summary_to_markdown(meeting))
    elif fmt == "docx":
        p = summary_to_docx(meeting, d / f"{meeting['id']}-summary.docx")
    else:
        raise ValueError(f"unknown format {fmt}")
    return p
